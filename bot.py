import telebot
from telebot import types
from docx import Document
import tempfile
import os
from datetime import datetime
from Messages_kz import *
from Messages_ru import *
from docx.shared import Pt
import calendar

Token = '6916931629:AAG0EB9xsAgJ5RLS_S3BErGwIhGQnr3pWJQ'
bot = telebot.TeleBot(Token)
CHAT_ID = 950968361
previous_messages = {}
messages = []
user_data = {}
question_stack = {}
user_language = {}
questions1 = {
    'Name_student': 'Введите ваше ФИО:',
    'student_group': 'Введите вашу группу:',
    'date_lesson': 'Введите дату занятия:',
    'prichina': 'Введите причину отсутствия на занятие',
}
questions1_kz = {
    'Name_student': 'Аты-жөніңізді енгізіңіз:',
    'student_gender': 'Сіз студент пе, не студентка ма?',
    'student_group': 'Топты енгізіңіз:',
    'date_lesson_kz': 'Сабақ күнін енгізіңіз:',
    'prichina': 'Сабаққа бармау үшін себепті енгізіңіз:',
}
questions2 = {
    'Student_name': 'ВВедите ваше ФИО',
    'Student_group': 'Введите вашу группу',
    'Diplom_ruk': 'Введите имя дипломного руководителя'
}
questions2_kz = {
    'Student_name': 'Аты-жөніңізді енгізіңіз',
    'Student_group': 'Топты енгізіңіз',
    'student_gender': 'Сіз студент пе, не студентка ма?',
    'Diplom_ruk': 'Дипломды басқарушының атын енгізіңіз',
}
questions3 = {
    'Student_name': 'Введите ваше ФИО',
    "nu": 'Введите номер приказа',
    "day": 'Введите день приказа',
    "mont": 'Введите месяц приказа',
    "year": 'Введите год прикаха',
    "course": 'Введите курс',
    "student_group": 'Введите группу студента',
    "prediavlent": 'Введите кому предьявляется'
}
questions3_kz = {
    'Student_name': 'Аты-жөніңізді енгізіңіз',
    "nu": 'Приказ нөмірін енгізіңіз',
    "day": 'Приказ күнін енгізіңіз',
    "mont": 'Приказ айын енгізіңіз',
    "year": 'Приказ жылын енгізіңіз',
    "course": 'Курсты енгізіңіз',
    "student_group": 'Студентті топты енгізіңіз',
    "prediavlent": 'Кімге көз көрсетілгенді енгізіңіз'

}
questions4 = {
    'Student_name': 'Введите свое ФИО',
    'stay_place': 'Введите ваще место проживание',
    'country': 'Введите страну',
    'Addres_life': 'Введите адресс проживания',
    'Addres_registration': 'Введите адресс регистрации',
    'School_end': 'Введите год окончание учебного заведения',
    'Doc_study': 'У вас аттестат или диплом?',
    'altyn_belgi': 'У вас есть алтын белги или вы закончили с отличием?',
    'Language_otdel': 'Введите на какое языковое отделение хотите',
    'inastr_language': 'Какой иностранный язык изучали',
    'level_of_language': 'Уровень иностранного языка',
    'rezhim_uchebi': 'Какой режим учебы хотите? (очное, дистанционное) ',
    'stepen_akadem': 'На какую академическую степень хотите поступить',
    'fakultet': 'Введите на какой факультет',
    'group_obrazovat_programm': 'Введите группу обрзовательных прорамм(4 значный код и название)',
    'obrazovat_programm': 'Введите образовательную программу',
    'type_obrazovaniya': 'Платное или грантное обучение?',
    'financirovanie': 'Введите вид финансирование',
    'Nomer_certificate': 'Номер сертификата ЕНТ или КТ',
    'nomer_granta': 'Введите номер гранта(в случии отсутствия напишите -)',
    'obshezhitiya': 'Вам нужно общежитие?',
    'phone': 'Ваш номер телефона',
    'Father': 'Введите информацию Отца в ввиде ФИО,телефон,адрес регистрации',
    'Mother': 'Введите информацию Матери в ввиде ФИО,телефон,адрес регистрации',
    'Ligot': 'У вас есть льготы?',

}
questions4_kz = {
    'Student_name': 'Есімінізді енгізіңіз',
    'stay_place': 'Тұрмысыңызды енгізіңіз',
    'country': 'Жататын өлкені енгізіңіз',
    'Addres_life': 'Тұрмысыңыздың мекенжайын енгізіңіз',
    'Addres_registration': 'Тіркеу көмегімен мекенжайын енгізіңіз',
    'School_end': 'Оқу орнын аяқтаган жылды енгізіңіз',
    'Doc_study': 'Сізде аттестат пен диплом бар ма?',
    'altyn_belgi': 'Сізде алтын белгісі немесе отличие бар ма?',
    'Language_otdel': 'Қай тіл бөліміне кіру келгенін енгізіңіз',
    'inastr_language': 'Қай иелі сауатты оқуды енгізіңіз',
    'level_of_language': 'Иелі сауатты деңгейі',
    'rezhim_uchebi': 'Оқу режимін таңдауыңызды енгізіңіз? (күнделікті, қашықтықты)',
    'stepen_akadem': 'Қай академиялық деңгейге кіру келгенін енгізіңіз',
    'fakultet': 'Қай факультетке кіру келгенін енгізіңіз',
    'group_obrazovat_programm': 'Тәлім беру бағдарламасының топтық коды мен атауын енгізіңіз (4 таңбалы код және атау)',
    'obrazovat_programm': 'Тәлім беру бағдарламасын енгізіңіз',
    'type_obrazovaniya': 'Тәлімді жаттығу немесе грант бойынша алу келгенін енгізіңіз?',
    'financirovanie': 'Қаржылау түрін енгізіңіз',
    'Nomer_certificate': 'ЕНТ немесе КТ сертификатының нөмірін енгізіңіз',
    'nomer_granta': 'Гранттың нөмірін енгізіңіз (болмаса, - деп жазыңыз)',
    'obshezhitiya': 'Сізге қолаймыз келетін ба? (жоқ, ия)',
    'phone': 'Сіздің телефон нөміріңіз',
    'Father': 'Әке туралы ақпаратты ФИО, телефон, тіркеу мекенжайы форматында енгізіңіз',
    'Mother': 'Ана туралы ақпаратты ФИО, телефон, тіркеу мекенжайы форматында енгізіңіз',
    'Ligot': 'Сізде льготалар бар ма?',
}

month_names_ru = {
    'January': 'января',
    'February': 'февраля',
    'March': 'марта',
    'April': 'апреля',
    'May': 'мая',
    'June': 'июня',
    'July': 'июля',
    'August': 'августа',
    'September': 'сентября',
    'October': 'октября',
    'November': 'ноября',
    'December': 'декабря',
}
month_names_kz = {
    'January': 'қаңтар',
    'February': 'ақпан',
    'March': 'наурыз',
    'April': 'сәуір',
    'May': 'мамыр',
    'June': 'маусым',
    'July': 'шілде',
    'August': 'тамыз',
    'September': 'қыркүйек',
    'October': 'қазан',
    'November': 'қараша',
    'December': 'желтоқсан'
}


def social_info(message):
    instagram_link = "https://www.instagram.com/aliakbarisainov/"
    vk_link = "https://vk.com/isain2000"
    social_media_info = f"Инстаграм: {instagram_link}\nВКонтакте: {vk_link}"
    bot.send_message(message.chat.id, social_media_info)


def send_error_message(chat_id, error_text):
    bot.send_message(chat_id, f" {error_text}")


def fill_document(template_path, data):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Time New Roman'
            run.font.size = Pt(14)

    for key, value in data.items():
        for paragraph in doc.paragraphs:
            if f'{{{key}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))
                for run in paragraph.runs:
                    run.font.size = Pt(14)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if f'{{{key}}}' in cell.text:
                        cell.text = cell.text.replace(f'{{{key}}}', str(value))
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(14)

    return doc


def send_document(chat_id, doc):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_file_name = temp_file.name
        doc.save(temp_file_name)
        temp_file.close()
        with open(temp_file_name, 'rb') as file:
            bot.send_document(chat_id, file)
        os.unlink(temp_file_name)


def ask_next_question(message, user_id, template_path, questions):
    if user_id not in user_data:
        user_data[user_id] = {}

    variables = list(questions.keys())
    current_index = len(user_data[user_id])

    if current_index < len(variables):
        next_variable = variables[current_index]
        question_text = questions[next_variable]
        bot.send_message(message.chat.id, question_text)

        user_data[user_id][next_variable] = None
        import functools
        callback = functools.partial(process_answer, user_id=user_id, variable=next_variable,
                                     template_path=template_path, questions=questions)

        bot.register_next_step_handler(message, callback)
    else:
        generate_and_send_document(message, user_id, template_path, questions)


def process_answer(message, user_id, variable, template_path, questions):
    answer = message.text
    if answer.lower() == '/stop':
        if user_language == 'ru':
            bot.send_message(message.chat.id, "Опрос прерван.")
        else:
            bot.send_message(message.chat.id, "Сауалнама үзілді")
        if user_id in user_data:
            del user_data[user_id]
            if user_language == 'kz':
                send_start_message_kz(message)
            else:
                send_start_message(message)
    else:
        if variable == 'date_lesson':
            current_datetime = datetime.now()
            try:
                if '-' in answer:
                    start_date_str, end_date_str = answer.split('-')
                    start_date = datetime.strptime(start_date_str.strip(), '%d.%m.%Y')
                    end_date = datetime.strptime(end_date_str.strip(), '%d.%m.%Y')

                    if start_date.year > current_datetime.year or start_date.year < (current_datetime.year - 1):
                        raise ValueError("Некорректный год")

                    if start_date.month > 12 or start_date.day > 31:
                        raise ValueError("Некорректный месяц или день")
                    if end_date.year > current_datetime.year or end_date.year < (current_datetime.year - 1):
                        raise ValueError("Некорректный год")

                    if end_date.month > 12 or end_date.day > 31:
                        raise ValueError("Некорректный месяц или день")
                    user_data[user_id][
                        variable] = f"от {start_date.strftime('%d.%m.%Y')} по {end_date.strftime('%d.%m.%Y')}"
                else:
                    date_obj = datetime.strptime(answer, '%d.%m.%Y')
                    if date_obj.year > current_datetime.year or date_obj.year < (current_datetime.year - 1):
                        raise ValueError("Некорректный год")
                    if date_obj.month > 12 or date_obj.day > 31:
                            raise ValueError("Некорректный месяц или день")
                    user_data[user_id][variable] = date_obj.strftime('%d.%m.%Y')
                ask_next_question(message, user_id, template_path, questions)
            except ValueError:
                    send_error_message(message.chat.id,
                                       'Некорректные даты. Пожалуйста, введите даты в формате DD.MM.YYYY и убедитесь в правильности значений.')
                    bot.send_message(message.chat.id, questions[variable])
                    bot.register_next_step_handler(message, process_answer, user_id=user_id, variable=variable,
                                               template_path=template_path, questions=questions)
            if variable == 'date_lesson_kz':
                current_datetime = datetime.now()
                try:
                    if '-' in answer:
                        start_date_str, end_date_str = answer.split('-')
                        start_date = datetime.strptime(start_date_str.strip(), '%d.%m.%Y')
                        end_date = datetime.strptime(end_date_str.strip(), '%d.%m.%Y')

                        if start_date.year > current_datetime.year or start_date.year < (current_datetime.year - 1):
                                raise ValueError("Дұрыс емес жыл")

                        if start_date.month > 12 or start_date.day > 31:
                                raise ValueError("Ай немесе күн жарамсыз")

                        if end_date.year > current_datetime.year or end_date.year < (current_datetime.year - 1):
                                raise ValueError("Дұрыс емес жыл")
                        if end_date.month > 12 or end_date.day > 31:
                                raise ValueError("Ай немесе күн жарамсыз")
                        user_data[user_id][
                            variable] = f"от {start_date.strftime('%d.%m.%Y')} по {end_date.strftime('%d.%m.%Y')}"
                    else:
                        date_obj = datetime.strptime(answer, '%d.%m.%Y')
                        if date_obj.year > current_datetime.year or date_obj.year < (current_datetime.year - 1):
                                raise ValueError("Дұрыс емес жыл")

                        if date_obj.month > 12 or date_obj.day > 31:
                                raise ValueError("Ай немесе күн жарамсыз")
                        user_data[user_id][variable] = date_obj.strftime('%d.%m.%Y')

                    ask_next_question(message, user_id, template_path, questions)
                except ValueError:
                        send_error_message(message.chat.id,
                                           'Дұрыс емес күндер. Күндерді DD.MM.YYYY форматында енгізіп, мәндердің дұрыстығына көз жеткізіңіз.'
                                           '')
                        bot.send_message(message.chat.id, questions[variable])
                        bot.register_next_step_handler(message, process_answer, user_id=user_id, variable=variable,
                                                   template_path=template_path, questions=questions)
        else:
            user_data[user_id][variable] = answer
            ask_next_question(message, user_id, template_path, questions)


def generate_and_send_document(message, user_id, template_path, questions):
    if user_id in user_data:
        document_data = user_data[user_id]
        current_datetime = datetime.now()
        current_month = current_datetime.month

        if current_month >= 9:  # September and onwards
            next_year = current_datetime.year + 1
            last_day = datetime(next_year, 6, calendar.monthrange(next_year, 6)[1])
        else:
            last_day = datetime(current_datetime.year, 6, calendar.monthrange(current_datetime.year, 6)[1])

        document_data['last_day'] = last_day.strftime('%d.%m.%Y')
        filled_document = fill_document(template_path, document_data)
        send_document(message.chat.id, filled_document)
        del user_data[user_id]


document_path = 'Docs/Courts/o_priznanii_grazhdanina_nedeesposobnym.docx'
document_path1 = 'Docs/Courts/o_vozmeshchenii_ushcherba_dtp.docx'
document_path2 = 'Docs/Courts/ob_ustanovlenii_fakta_smerti_grazhdanina.docx'


def send_document_asc(chat_id):
    try:
        with open(document_path, 'rb') as document:
            bot.send_document(chat_id, document)
    except Exception as e:
        print(f"Ошибка при отправке документа: {e}")


def send_document_acs(chat_id):
    try:
        with open(document_path1, 'rb') as document:
            bot.send_document(chat_id, document)
    except Exception as e:
        print(f"Ошибка при отправке документа: {e}")


def send_document_abs(chat_id):
    try:
        with open(document_path2, 'rb') as document:
            bot.send_document(chat_id, document)
    except Exception as e:
        print(f"Ошибка при отправке документа: {e}")


def send_start_message(message):
    markup = types.InlineKeyboardMarkup(row_width=1)
    item1 = types.InlineKeyboardButton('АТБ', callback_data='univer')
    item2 = types.InlineKeyboardButton("Колледж", callback_data='college')
    item3 = types.InlineKeyboardButton("Правохранительные органы", callback_data='courts')
    item4 = types.InlineKeyboardButton('Веб-сайт колледжа', url='https://polytech.kz/')
    markup.add(item1, item2, item3,item4)
    sent_message = bot.send_message(message.chat.id, "Выберите образец который хотите получить", reply_markup=markup)
    previous_messages[message.chat.id] = [sent_message.message_id]


def send_start_message_kz(message):
    markup = types.InlineKeyboardMarkup(row_width=1)
    item1 = types.InlineKeyboardButton('АТБ', callback_data='univer_kz')
    item2 = types.InlineKeyboardButton("Колледж", callback_data='college_kz')
    item3 = types.InlineKeyboardButton("Құқықты қауіпсіздік орталықтар", callback_data='courts_kz')
    item4 = types.InlineKeyboardButton('Колледждің веб-сайты', url='https://polytech.kz/')
    markup.add(item1, item2, item3,item4)
    sent_message = bot.send_message(message.chat.id, "Образецті таңдау", reply_markup=markup)
    previous_messages[message.chat.id] = [sent_message.message_id]


def language_set(message):
    language_markup = types.InlineKeyboardMarkup(row_width=3)
    item1 = types.InlineKeyboardButton("kz Казахский", callback_data='kz')
    item2 = types.InlineKeyboardButton("ru Русский", callback_data='ru')
    language_markup.add(item1, item2)
    sent_message = bot.send_message(message.chat.id, "Выберите язык", reply_markup=language_markup)


@bot.message_handler(commands=['start'])
def start_message(message):
    language_set(message)


@bot.message_handler(commands=['language'])
def language_pick(message):
    language_set(message)

@bot.message_handler(commands=['help'])
def help(message):
    start = bot.send_message(message.chat.id, 'пожалуйста опишите свою проблему или ошибку')
    bot.register_next_step_handler(start,bugs_report)
def bugs_report(message):
    user_id = message.from_user.id
    user_name = message.from_user.username
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    message_to_save = message.text
    bot.send_message(message.chat.id,'Ваш запрос отправлена админу')
    bot.send_message(chat_id=CHAT_ID,
                     text=f"Пользователь {user_id} ({user_name}) отправил вопрос/ошибку:\n\n{message_to_save}\n\nДата: {current_time}")
@bot.message_handler(commands=['info'])
def info(message):
    if user_language == 'kz':
        user_language[message.chat.id] = 'kz'
        bot.send_message(message.chat.id, f"Сіз тілді таңдағаныз:Казахский")
        bot.delete_message(message.chat.id, message.message_id)
        bot.send_message(message.chat.id, start_kz)
        social_info(message)
        bot.send_message(message.chat.id, 'Тіл ауыстыру үшін /language деп жазыныз')
        bot.send_message(message.text, 'Егер сізде қателіктер немесе сұраулар туындаса, /help командасын енгізіп жазыңдарсыз.')
        bot.send_message(message.chat.id, 'Бастау үшін Құжаттар деп жазыныз')
    elif user_language == 'ru':
        user_language[message.chat.id] = 'ru'
        bot.send_message(message.chat.id, f"Вы выбрали язык:Русский")
        bot.delete_message(message.chat.id, message.message_id)
        bot.send_message(message.chat.id, start_ru)
        social_info(message)
        bot.send_message(message.text, 'Для смены языка напишите команду /language')
        bot.send_message(message.text, 'Если у вас возникли ошибки или вопросы, напишите команду /help')
        bot.send_message(message.chat.id, 'Для начала напишите Документы')
    else:
        bot.send_message(message.chat.id,
                         'Вы не выбрали язык, пожалуйста выберите язык. Для выбора языка напишите /start и выберите ваш язык')
        bot.send_message(message.chat.id,
                         'Сіз тілді таңдауды алмадыңыз, өтініш тілді таңдаңыз. Тілді тандау үшін /start жазып өзініздің тіліңізді танданыз ')


@bot.message_handler(func=lambda message: True)
def handle_message(message):
    if message.text.lower() == 'документы' or message.text.upper == 'Документы':
        send_start_message(message)
    elif message.text.lower() == 'құжаттар' or message.text.upper() == 'Құжаттар':
        send_start_message_kz(message)
    elif not user_language:
        bot.send_message(message.chat.id,
                         'Вы не выбрали язык, пожалуйста выберите язык. Для выбора языка напишите /start и выберите ваш язык')
        bot.send_message(message.chat.id,
                         'Сіз тілді таңдауды алмадыңыз, өтініш тілді таңдаңыз. Тілді тандау үшін /start жазып өзініздің тіліңізді танданыз ')
    else:
        bot.send_message(message.chat.id, "Дұрыс емес болмаған команда. Мәлімет алу үшін /info жазыңыз.")
        bot.send_message(message.chat.id, "Неправильная команда. Напишите /info для получения информации.")


@bot.callback_query_handler(func=lambda call: True)
def handle_callback_query(call):
    user_id = call.from_user.id
    if call.message.chat.id in previous_messages:
        for msg_id in previous_messages[call.message.chat.id]:
            bot.delete_message(call.message.chat.id, msg_id)
        previous_messages[call.message.chat.id] = []

    if call.data == 'kz':
        user_language[call.message.chat.id] = 'kz'
        bot.send_message(call.message.chat.id, f"Сіз тілді таңдағаныз:Казахский")
        bot.delete_message(call.message.chat.id, call.message.message_id)
        bot.send_message(call.message.chat.id, start_kz)
        social_info(call.message)
        bot.send_message(call.message.chat.id, 'Тіл ауыстыру үшін /language деп жазыныз')
        bot.send_message(call.message.chat.id, 'Бастау үшін Құжаттар деп жазыныз')
    elif call.data == 'ru':
        user_language[call.message.chat.id] = 'ru'
        bot.send_message(call.message.chat.id, f"Вы выбрали язык:Русский")
        bot.delete_message(call.message.chat.id, call.message.message_id)
        bot.send_message(call.message.chat.id, start_ru)
        social_info(call.message)
        bot.send_message(call.message.chat.id, 'Для начала напишите Документы')
        bot.send_message(call.message.chat.id, 'Для смены языка напишите команду /language')

    if call.data == 'college_button1':

        bot.send_message(call.message.chat.id, "Начнем опрос. Ответьте на следующие вопросы:")
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='PHOTOS/College/Обьяснительная.docx',
                          questions=questions1)
    elif call.data == 'college_button1_kz':
        bot.send_message(call.message.chat.id, "Сауалнаманы бастау. Алдағы сұрауларға жауап беріңіз")
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='PHOTOS/College/Обьяснительная_kz.docx',
                          questions=questions1_kz)
    elif call.data == 'college_button2':
        bot.send_message(call.message.chat.id, "Начнем опрос. Ответьте на следующие вопросы:")
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='PHOTOS/College/Заявление.docx',
                          questions=questions2)
    elif call.data == 'college_button2_kz':
        bot.send_message(call.message.chat.id, "Сауалнаманы бастау. Алдағы сұрауларға жауап беріңіз")
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='PHOTOS/College/Заявление_kz.docx',
                          questions=questions2_kz)
    elif call.data == 'college_button3':
        bot.send_message(call.message.chat.id, "Начнем опрос. Ответьте на следующие вопросы:")
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='PHOTOS/College/СПРАВКА.docx', questions=questions3)
    elif call.data == 'college_button3_kz':
        bot.send_message(call.message.chat.id, "Сауалнаманы бастау. Алдағы сұрауларға жауап беріңіз")
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='PHOTOS/College/СПРАВКА.docx',
                          questions=questions3_kz)
    elif call.data == 'college':
        markup2 = types.InlineKeyboardMarkup(row_width=2)
        item3_1 = types.InlineKeyboardButton("Объяснительная", callback_data='college_button1')
        item3_2 = types.InlineKeyboardButton("Заявления на дипломного руководителя ",
                                             callback_data='college_button2')
        item3_3 = types.InlineKeyboardButton("Справка", callback_data='college_button3')
        item3_4 = types.InlineKeyboardButton("Назад", callback_data='college_back_button')
        markup2.add(item3_1, item3_2, item3_3, item3_4)
        sent_message = bot.send_message(chat_id=call.message.chat.id, text="Выберите документ",
                                        reply_markup=markup2)
        previous_messages[call.message.chat.id] = [sent_message.message_id]
    elif call.data == 'college_kz':
        markup1 = types.InlineKeyboardMarkup(row_width=2)
        item2_1 = types.InlineKeyboardButton("Түсіндіруші хат", callback_data='college_button1_kz')
        item2_2 = types.InlineKeyboardButton("Дипломды басқарушыға өтініштер ",
                                             callback_data='college_button2_kz')
        item2_3 = types.InlineKeyboardButton("Сұрау", callback_data='college_button3_kz')
        item2_4 = types.InlineKeyboardButton("Кері қайту", callback_data='college_back_button_kz')
        markup1.add(item2_1, item2_2, item2_3, item2_4)
        sent_message = bot.send_message(chat_id=call.message.chat.id, text="Құжатты таңдау",
                                        reply_markup=markup1)
        previous_messages[call.message.chat.id] = [sent_message.message_id]
    elif call.data == 'college_back_button_kz':
        send_start_message_kz(call.message)
    elif call.data == 'college_back_button':
        send_start_message(call.message)
    if call.data == 'courts':
        markup1 = types.InlineKeyboardMarkup(row_width=2)
        item3_1 = types.InlineKeyboardButton("Признание гражданина недееспособным", callback_data='courts_button1')
        item3_2 = types.InlineKeyboardButton("О возмещение ущерба дтп", callback_data='courts_button2')
        item3_3 = types.InlineKeyboardButton("Об установление факта смерти", callback_data='courts_button3')
        item3_4 = types.InlineKeyboardButton("Назад", callback_data='courts_back_button')
        markup1.add(item3_1, item3_2, item3_3, item3_4)
        sent_message = bot.send_message(chat_id=call.message.chat.id, text="Выберите документ",
                                        reply_markup=markup1)
        previous_messages[call.message.chat.id] = [sent_message.message_id]
    elif call.data == 'courts_kz':
        markup1 = types.InlineKeyboardMarkup(row_width=2)
        item3_1 = types.InlineKeyboardButton("Гражданын елестік емес деп тану", callback_data='courts_button1')
        item3_2 = types.InlineKeyboardButton("Автошаға туынды дамуды компенсациялау",
                                             callback_data='courts_button2')
        item3_3 = types.InlineKeyboardButton("Құжатты таңдау", callback_data='courts_button3')
        item3_4 = types.InlineKeyboardButton("Кері қайту", callback_data='courts_back_button_kz')
        markup1.add(item3_1, item3_2, item3_3, item3_4)
        sent_message = bot.send_message(chat_id=call.message.chat.id, text="Құжатты таңдау",
                                        reply_markup=markup1)
        previous_messages[call.message.chat.id] = [sent_message.message_id]
    elif call.data == 'courts_back_button_kz':
        send_start_message_kz(call.message)
    elif call.data == 'courts_back_button':
        send_start_message(call.message)
    if call.data == 'courts_button1':
        send_document_asc(user_id)
    if call.data == 'courts_button2':
        send_document_acs(user_id)
    if call.data == 'courts_button3':
        send_document_abs(user_id)
    if call.data == 'univer_kz':
        markup1 = types.InlineKeyboardMarkup(row_width=1)
        item1 = types.InlineKeyboardButton('Сұрау', callback_data='univer_button1_kz')
        markup1.add(item1)
        sent_message = bot.send_message(chat_id=call.message.chat.id, text="Құжатты таңдау",
                                        reply_markup=markup1)
        previous_messages[call.message.chat.id] = [sent_message.message_id]
    if call.data == 'univer':
        markup1 = types.InlineKeyboardMarkup(row_width=1)
        item1 = types.InlineKeyboardButton('Заявление', callback_data='univer_button1')
        markup1.add(item1)
        sent_message = bot.send_message(chat_id=call.message.chat.id, text="Выберите документ",
                                        reply_markup=markup1)
        previous_messages[call.message.chat.id] = [sent_message.message_id]
    if call.data == 'univer_button1':
        bot.send_message(call.message.chat.id, "Начнем опрос. Ответьте на следующие вопросы:")
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='Docs/Universet/Заявление.docx', questions=questions4)
    if call.data == 'univer_button1_kz':
        bot.send_message(call.message.chat.id, "Сауалнаманы бастау. Алдағы сұрауларға жауап беріңіз")
        bot.send_message(call.message.chat.id, 'Тоқтату үшін /stop деп жазыныз')
        user_data[user_id] = {}
        ask_next_question(call.message, user_id, template_path='Docs/Universet/Заявление_kz.docx',
                          questions=questions4_kz)


bot.polling(none_stop=True, timeout=123)
